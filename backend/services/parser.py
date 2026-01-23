"""
Document Parser Service.
Parses DOCX files and returns structured JSON data.
Based on view_docx.py with JSON output.
"""
import docx
from docx.shared import RGBColor
from typing import Optional
import io


def rgb_to_hex(color: Optional[RGBColor]) -> Optional[str]:
    """Convert RGBColor to hex string."""
    if color is None:
        return None
    return f"#{color.red:02x}{color.green:02x}{color.blue:02x}"


def get_run_info(run) -> dict:
    """Extract formatting info from a run."""
    color = None
    if run.font.color and run.font.color.rgb:
        color = rgb_to_hex(run.font.color.rgb)
    
    font_size = None
    if run.font.size:
        font_size = run.font.size.pt
    
    return {
        "text": run.text,
        "bold": bool(run.bold),
        "italic": bool(run.italic),
        "underline": bool(run.underline),
        "color": color,
        "font_size": font_size
    }


def parse_document(file_content: bytes) -> dict:
    """
    Parse a DOCX file and return structured data.
    
    Args:
        file_content: Raw bytes of the DOCX file
        
    Returns:
        Dictionary with paragraphs, tables, stats, and HTML preview
    """
    doc = docx.Document(io.BytesIO(file_content))
    
    paragraphs = []
    styles = set()
    
    # Parse body paragraphs
    for i, para in enumerate(doc.paragraphs):
        style_name = para.style.name if para.style else "Normal"
        styles.add(style_name)
        
        runs = [get_run_info(run) for run in para.runs]
        
        paragraphs.append({
            "index": i,
            "text": para.text,
            "style": style_name,
            "runs": runs
        })
    
    # Parse tables
    tables_count = len(doc.tables)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.style:
                        styles.add(para.style.name)
    
    # Generate HTML preview
    html_preview = generate_html_preview(doc)
    
    return {
        "paragraphs": paragraphs,
        "stats": {
            "total_paragraphs": len(paragraphs),
            "total_tables": tables_count,
            "unique_styles": sorted(list(styles))
        },
        "html_preview": html_preview
    }


def generate_html_preview(doc) -> str:
    """Generate HTML preview of the document."""
    html_parts = []
    
    for para in doc.paragraphs:
        if not para.text.strip():
            continue
            
        style = para.style.name if para.style else "Normal"
        
        # Map DOCX styles to HTML tags
        if "Heading 1" in style:
            tag = "h1"
        elif "Heading 2" in style:
            tag = "h2"
        elif "Heading 3" in style:
            tag = "h3"
        elif "Title" in style:
            tag = "h1"
        else:
            tag = "p"
        
        # Build content with run formatting
        content_parts = []
        for run in para.runs:
            text = run.text
            if not text:
                continue
                
            # Apply inline styles
            style_attrs = []
            if run.font.color and run.font.color.rgb:
                style_attrs.append(f"color:{rgb_to_hex(run.font.color.rgb)}")
            
            if run.bold:
                text = f"<strong>{text}</strong>"
            if run.italic:
                text = f"<em>{text}</em>"
            if run.underline:
                text = f"<u>{text}</u>"
            
            if style_attrs:
                text = f'<span style="{";".join(style_attrs)}">{text}</span>'
            
            content_parts.append(text)
        
        content = "".join(content_parts)
        if content:
            html_parts.append(f"<{tag}>{content}</{tag}>")
    
    return "\n".join(html_parts)
