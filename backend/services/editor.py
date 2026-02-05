"""
Document Editor Service.
Handles text replacement while preserving styles.
Based on modify_docx.py with enhanced features.
"""
import docx
import io
import re
from typing import Optional
from .parser import generate_html_preview


def replace_text_in_run(run, old_text: str, new_text: str, case_sensitive: bool = False) -> bool:
    """
    Replace text in a run while preserving formatting.
    
    Args:
        run: A docx Run object
        old_text: Text to find
        new_text: Text to replace with
        case_sensitive: Whether to match case
        
    Returns:
        True if replacement was made
    """
    if case_sensitive:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)
            return True
    else:
        pattern = re.compile(re.escape(old_text), re.IGNORECASE)
        if pattern.search(run.text):
            run.text = pattern.sub(new_text, run.text)
            return True
    return False


def replace_in_paragraph(paragraph, old_text: str, new_text: str, 
                         case_sensitive: bool = False) -> bool:
    """
    Replace text in all runs of a paragraph.
    
    Returns:
        True if any replacement was made
    """
    replaced = False
    for run in paragraph.runs:
        if replace_text_in_run(run, old_text, new_text, case_sensitive):
            replaced = True

    if replaced:
        return True

    # Fallback: replace across run boundaries (may lose inline formatting in this paragraph)
    if paragraph.text:
        if case_sensitive:
            if old_text in paragraph.text:
                paragraph.text = paragraph.text.replace(old_text, new_text)
                return True
        else:
            pattern = re.compile(re.escape(old_text), re.IGNORECASE)
            if pattern.search(paragraph.text):
                paragraph.text = pattern.sub(new_text, paragraph.text)
                return True

    return False


def replace_in_document(
    doc,
    replacements: list[dict],
    case_sensitive: bool = False
) -> dict:
    """
    Apply all replacements to a document.
    
    Args:
        doc: python-docx Document object
        replacements: List of {"find": str, "replace": str}
        case_sensitive: Whether to match case
        
    Returns:
        Dictionary with replacement statistics
    """
    total = 0
    affected = set()
    
    # Process body paragraphs
    for i, para in enumerate(doc.paragraphs):
        for rep in replacements:
            if replace_in_paragraph(para, rep["find"], rep["replace"], case_sensitive):
                total += 1
                affected.add(i)
    
    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for rep in replacements:
                        if replace_in_paragraph(
                            para, rep["find"], rep["replace"], case_sensitive
                        ):
                            total += 1
    
    # Process headers and footers
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header:
                for para in header.paragraphs:
                    for rep in replacements:
                        replace_in_paragraph(
                            para, rep["find"], rep["replace"], case_sensitive
                        )
                for table in header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                for rep in replacements:
                                    replace_in_paragraph(
                                        para, rep["find"], rep["replace"], case_sensitive
                                    )
        
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer:
                for para in footer.paragraphs:
                    for rep in replacements:
                        replace_in_paragraph(
                            para, rep["find"], rep["replace"], case_sensitive
                        )
                for table in footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                for rep in replacements:
                                    replace_in_paragraph(
                                        para, rep["find"], rep["replace"], case_sensitive
                                    )
    
    return {
        "total_replacements": total,
        "affected_paragraphs": sorted(list(affected)),
        "html_preview": generate_html_preview(doc)
    }
