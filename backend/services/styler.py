"""
Style Modification Service.
Handles changing styles (colors, bold, italic) based on matching criteria.
"""
import docx
from docx.shared import RGBColor
from typing import Optional
from .parser import rgb_to_hex, generate_html_preview
from docx.oxml.ns import qn


COLOR_ALIASES = {
    "purple": "#800080",
    "morado": "#800080",
    "violet": "#8a2be2",
    "magenta": "#ff00ff",
    "red": "#ff0000",
    "rojo": "#ff0000",
    "blue": "#0000ff",
    "azul": "#0000ff",
    "green": "#00ff00",
    "verde": "#00ff00",
    "black": "#000000",
    "negro": "#000000",
    "white": "#ffffff",
    "blanco": "#ffffff",
}

COLOR_GROUPS = {
    "purple": (250, 330),
    "morado": (250, 330),
    "violet": (250, 300),
    "magenta": (300, 340),
    "red": (340, 20),
    "rojo": (340, 20),
    "blue": (200, 260),
    "azul": (200, 260),
    "green": (90, 150),
    "verde": (90, 150),
}

HIGHLIGHT_PURPLE_NAMES = {"purple", "magenta", "violet", "darkMagenta"}
HIGHLIGHT_RED_NAMES = {"red", "darkRed"}
HIGHLIGHT_BLUE_NAMES = {"blue", "darkBlue"}
HIGHLIGHT_GREEN_NAMES = {"green", "darkGreen"}
HIGHLIGHT_GRAY_NAMES = {"lightGray", "darkGray", "black"}

HIGHLIGHT_GROUPS = {
    "purple": HIGHLIGHT_PURPLE_NAMES,
    "morado": HIGHLIGHT_PURPLE_NAMES,
    "violet": HIGHLIGHT_PURPLE_NAMES,
    "magenta": HIGHLIGHT_PURPLE_NAMES,
    "red": HIGHLIGHT_RED_NAMES,
    "rojo": HIGHLIGHT_RED_NAMES,
    "blue": HIGHLIGHT_BLUE_NAMES,
    "azul": HIGHLIGHT_BLUE_NAMES,
    "green": HIGHLIGHT_GREEN_NAMES,
    "verde": HIGHLIGHT_GREEN_NAMES,
}

HIGHLIGHT_NAME_SET = (
    HIGHLIGHT_PURPLE_NAMES
    | HIGHLIGHT_RED_NAMES
    | HIGHLIGHT_BLUE_NAMES
    | HIGHLIGHT_GREEN_NAMES
    | {"yellow", "darkYellow", "cyan", "darkCyan", "lightGray", "darkGray", "black", "none"}
)

GROUP_HUE_RANGES = {
    "red": (340, 20),
    "magenta": (300, 340),
    "purple": (250, 330),
    "blue": (200, 260),
    "green": (90, 150),
}

GROUP_TO_HIGHLIGHT = {
    "red": "red",
    "rojo": "red",
    "magenta": "magenta",
    "purple": "magenta",
    "morado": "magenta",
    "violet": "magenta",
    "blue": "blue",
    "azul": "blue",
    "green": "green",
    "verde": "green",
}


def normalize_color(value: Optional[str]) -> Optional[str]:
    if not value:
        return None
    raw = value.strip().lower()
    if raw in COLOR_ALIASES:
        return COLOR_ALIASES[raw]
    if raw.startswith('#') and len(raw) == 7:
        return raw
    if len(raw) == 6 and all(c in "0123456789abcdef" for c in raw):
        return f"#{raw}"
    return value


def hex_to_rgb_tuple(hex_color: str) -> Optional[tuple[int, int, int]]:
    if not hex_color:
        return None
    raw = hex_color.lstrip('#')
    if len(raw) != 6:
        return None
    try:
        return (int(raw[0:2], 16), int(raw[2:4], 16), int(raw[4:6], 16))
    except ValueError:
        return None


def rgb_to_hue_deg(r: int, g: int, b: int) -> Optional[float]:
    r_f, g_f, b_f = r / 255.0, g / 255.0, b / 255.0
    mx = max(r_f, g_f, b_f)
    mn = min(r_f, g_f, b_f)
    diff = mx - mn
    if diff == 0:
        return None
    if mx == r_f:
        h = (60 * ((g_f - b_f) / diff) + 360) % 360
    elif mx == g_f:
        h = (60 * ((b_f - r_f) / diff) + 120) % 360
    else:
        h = (60 * ((r_f - g_f) / diff) + 240) % 360
    return h


def get_color_group(raw_value: str) -> Optional[str]:
    if not raw_value:
        return None
    raw = raw_value.strip().lower()
    if raw in COLOR_GROUPS:
        return raw

    normalized = normalize_color(raw_value)
    if not normalized:
        return None

    rgb = hex_to_rgb_tuple(normalized)
    if not rgb:
        return None

    hue = rgb_to_hue_deg(*rgb)
    if hue is None:
        return None

    for group, (start, end) in GROUP_HUE_RANGES.items():
        if in_hue_range(hue, start, end):
            return group

    return None


def resolve_highlight_name(color_value: str) -> Optional[str]:
    if not color_value:
        return None
    raw = color_value.strip().lower()
    if raw in HIGHLIGHT_NAME_SET:
        return raw

    group = get_color_group(color_value)
    if group and group in GROUP_TO_HIGHLIGHT:
        return GROUP_TO_HIGHLIGHT[group]

    normalized = normalize_color(color_value)
    rgb = hex_to_rgb_tuple(normalized) if normalized else None
    if rgb:
        r, g, b = rgb
        if r == g == b:
            if r < 64:
                return "black"
            if r < 160:
                return "darkGray"
            return "lightGray"

    return None


def in_hue_range(hue: float, start: float, end: float) -> bool:
    if start <= end:
        return start <= hue <= end
    return hue >= start or hue <= end


def is_match_color_hex(hex_value: str, match_raw: str) -> bool:
    raw = match_raw.strip().lower()
    if raw in COLOR_GROUPS:
        rgb = hex_to_rgb_tuple(hex_value)
        if not rgb:
            return False
        hue = rgb_to_hue_deg(*rgb)
        if hue is None:
            return False
        start, end = COLOR_GROUPS[raw]
        return in_hue_range(hue, start, end)

    normalized = normalize_color(match_raw)
    if not normalized:
        return False
    return hex_value.lower() == normalized.lower()


def update_xml_colors(doc, match_color: str, apply_color: str) -> int:
    """Update colors in raw XML (font color, shading, highlight)."""
    changes = 0
    w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    val_attr = f"{{{w_ns}}}val"
    fill_attr = f"{{{w_ns}}}fill"

    apply_hex = normalize_color(apply_color) or "#ff0000"
    apply_hex = apply_hex.lower().lstrip('#')
    apply_val_hex = f"#{apply_hex}"
    apply_highlight = resolve_highlight_name(apply_color)

    for elem in doc._element.iter():
        tag = elem.tag.rsplit('}', 1)[-1]

        if tag == "color":
            val = elem.get(val_attr) or elem.get("val")
            if val and is_match_color_hex(val if val.startswith('#') else f"#{val}", match_color):
                elem.set(val_attr, apply_hex)
                changes += 1
        elif tag == "shd":
            fill = elem.get(fill_attr) or elem.get("fill")
            if fill and is_match_color_hex(fill if fill.startswith('#') else f"#{fill}", match_color):
                elem.set(fill_attr, apply_hex)
                changes += 1
        elif tag == "highlight":
            val = (elem.get(val_attr) or elem.get("val") or "").strip()
            if not val:
                continue
            val_lower = val.lower()
            match_raw = match_color.strip().lower()
            match_group = get_color_group(match_color)
            highlight_group = HIGHLIGHT_GROUPS.get(match_raw) or HIGHLIGHT_GROUPS.get(match_group)
            if highlight_group and val_lower in {v.lower() for v in highlight_group}:
                if apply_highlight:
                    elem.set(val_attr, apply_highlight)
                    changes += 1
            elif match_raw in HIGHLIGHT_NAME_SET and val_lower == match_raw:
                if apply_highlight:
                    elem.set(val_attr, apply_highlight)
                    changes += 1
            elif match_raw == val_lower:
                if apply_highlight:
                    elem.set(val_attr, apply_highlight)
                    changes += 1

    return changes


def hex_to_rgb(hex_color: str) -> RGBColor:
    """Convert hex color string to RGBColor."""
    hex_color = hex_color.lstrip('#')
    return RGBColor(
        int(hex_color[0:2], 16),
        int(hex_color[2:4], 16),
        int(hex_color[4:6], 16)
    )


def run_matches_criteria(run, match: dict, para_style: str) -> bool:
    """
    Check if a run matches the given criteria.
    
    Args:
        run: A docx Run object
        match: Dictionary with matching criteria (bold, italic, color, style)
        para_style: The paragraph's style name
        
    Returns:
        True if all specified criteria match
    """
    # Check style (paragraph level)
    if match.get("style") is not None:
        if match["style"].lower() not in para_style.lower():
            return False

    # Check text substring
    if match.get("text") is not None:
        if match["text"].lower() not in (run.text or "").lower():
            return False
    
    # Check bold
    if match.get("bold") is not None:
        if bool(run.bold) != match["bold"]:
            return False
    
    # Check italic
    if match.get("italic") is not None:
        if bool(run.italic) != match["italic"]:
            return False
    
    # Check color
    if match.get("color") is not None:
        run_color = None
        if run.font.color and run.font.color.rgb:
            run_color = rgb_to_hex(run.font.color.rgb)

        match_raw = (match.get("color") or "").strip().lower()
        if match_raw in COLOR_GROUPS:
            rgb = hex_to_rgb_tuple(run_color) if run_color else None
            if not rgb:
                return False
            hue = rgb_to_hue_deg(*rgb)
            if hue is None:
                return False
            start, end = COLOR_GROUPS[match_raw]
            if not in_hue_range(hue, start, end):
                return False
        else:
            match_color = normalize_color(match.get("color"))
            if run_color is None or not match_color or run_color.lower() != match_color.lower():
                return False
    
    return True


def apply_style_to_run(run, apply: dict) -> bool:
    """
    Apply style changes to a run.
    
    Args:
        run: A docx Run object
        apply: Dictionary with styles to apply (bold, italic, underline, color)
        
    Returns:
        True if any change was made
    """
    changed = False
    
    if apply.get("bold") is not None:
        run.bold = apply["bold"]
        changed = True
    
    if apply.get("italic") is not None:
        run.italic = apply["italic"]
        changed = True
    
    if apply.get("underline") is not None:
        run.underline = apply["underline"]
        changed = True
    
    if apply.get("color") is not None:
        normalized = normalize_color(apply["color"])
        if normalized:
            run.font.color.rgb = hex_to_rgb(normalized)
            changed = True
    
    return changed


def apply_style_changes(doc, changes: list[dict]) -> dict:
    """
    Apply style changes to the document.
    
    Args:
        doc: python-docx Document object
        changes: List of {"match": {...}, "apply": {...}}
        
    Returns:
        Dictionary with change statistics
    """
    total = 0
    affected = set()
    
    # Process body paragraphs
    for i, para in enumerate(doc.paragraphs):
        para_style = para.style.name if para.style else "Normal"
        
        for run in para.runs:
            if not run.text.strip():
                continue
                
            for change in changes:
                if run_matches_criteria(run, change["match"], para_style):
                    if apply_style_to_run(run, change["apply"]):
                        total += 1
                        affected.add(i)
    
    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para_style = para.style.name if para.style else "Normal"
                    for run in para.runs:
                        if not run.text.strip():
                            continue
                        for change in changes:
                            if run_matches_criteria(run, change["match"], para_style):
                                if apply_style_to_run(run, change["apply"]):
                                    total += 1
    
    # Update XML-level colors (shading/highlight/theme-level colors)
    for change in changes:
        match_color = change.get("match", {}).get("color")
        apply_color = change.get("apply", {}).get("color")
        if match_color and apply_color:
            total += update_xml_colors(doc, match_color, apply_color)

    return {
        "total_changes": total,
        "affected_paragraphs": sorted(list(affected)),
        "html_preview": generate_html_preview(doc)
    }
